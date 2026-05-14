import docx
from pathlib import Path
from .base import BaseParser
from ingestion.schema import RawDocument

class DOCXParser(BaseParser):
    def can_parse(self, file_path: str) -> bool:
        return Path(file_path).suffix.lower() in (".docx", ".doc")

    def parse(self, file_path: str) -> RawDocument:
        doc = docx.Document(file_path)
        text = "\n\n".join(p.text for p in doc.paragraphs if p.text.strip())
        return RawDocument(
            source_path=file_path,
            content_type="docx",
            raw_text=text,
        )